"""资料解析与标题优先、Token 兜底的 RAG 切片。"""

from pathlib import Path
import re

import tiktoken

from app.core.exceptions import RagConfigurationError

from .models import KnowledgeChunk, KnowledgeDocument


class KnowledgeDocumentParser:
    def parse_file(
        self, path: Path, *, knowledge_base_id: str, document_id: str
    ) -> KnowledgeDocument:
        if not path.is_file():
            raise RagConfigurationError(f"RAG 资料不存在: {path.name}")
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            content = path.read_text(encoding="utf-8")
        elif suffix == ".pdf":
            from pypdf import PdfReader

            content = "\n".join(
                page.extract_text() or "" for page in PdfReader(path).pages
            )
        elif suffix == ".docx":
            from docx import Document

            content = "\n".join(
                paragraph.text for paragraph in Document(path).paragraphs
            )
        else:
            raise RagConfigurationError(f"不支持的 RAG 资料格式: {suffix}")
        return KnowledgeDocument(
            knowledge_base_id=knowledge_base_id,
            document_id=document_id,
            source_name=path.name,
            content=content.strip(),
        )


class TokenChunker:
    _HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+?)\s*$")

    def __init__(self, *, chunk_size_tokens: int, overlap_tokens: int) -> None:
        if chunk_size_tokens < 1 or overlap_tokens < 0 or overlap_tokens >= chunk_size_tokens:
            raise RagConfigurationError("切片重叠 Token 必须小于切片大小")
        self._encoding = tiktoken.get_encoding("cl100k_base")
        self._chunk_size = chunk_size_tokens
        self._overlap_tokens = overlap_tokens

    def split(self, document: KnowledgeDocument) -> list[KnowledgeChunk]:
        sections = self._split_by_headings(document.content)
        chunks: list[KnowledgeChunk] = []
        for heading_path, heading_level, section_content in sections:
            for content, section_part_index in self._split_section(
                heading_path, section_content
            ):
                chunks.append(
                    KnowledgeChunk(
                        chunk_id=f"{document.document_id}:{len(chunks)}",
                        knowledge_base_id=document.knowledge_base_id,
                        document_id=document.document_id,
                        source_name=document.source_name,
                        chunk_index=len(chunks),
                        content=content,
                        metadata={
                            **document.metadata,
                            "chunkingStrategy": "heading_then_token",
                            "headingPath": " > ".join(item[1] for item in heading_path),
                            "headingLevel": str(heading_level),
                            "sectionPartIndex": str(section_part_index),
                        },
                    )
                )
        return chunks

    def _split_by_headings(
        self, content: str
    ) -> list[tuple[list[tuple[int, str]], int, str]]:
        """Return nonempty text sections, each associated with its Markdown heading path."""
        heading_path: list[tuple[int, str]] = []
        current_lines: list[str] = []
        sections: list[tuple[list[tuple[int, str]], int, str]] = []
        in_fenced_code_block = False

        def append_current() -> None:
            text = "\n".join(current_lines).strip()
            if text:
                sections.append((heading_path.copy(), heading_path[-1][0] if heading_path else 0, text))

        for line in content.splitlines():
            if line.lstrip().startswith("```"):
                current_lines.append(line)
                in_fenced_code_block = not in_fenced_code_block
                continue
            if in_fenced_code_block:
                current_lines.append(line)
                continue
            match = self._HEADING_PATTERN.match(line)
            if not match:
                current_lines.append(line)
                continue
            append_current()
            current_lines = []
            level = len(match.group(1))
            title = match.group(2).strip()
            heading_path = [item for item in heading_path if item[0] < level]
            heading_path.append((level, title))
        append_current()

        if sections:
            return sections
        return [([], 0, content.strip())] if content.strip() else []

    def _split_section(
        self, heading_path: list[tuple[int, str]], section_content: str
    ) -> list[tuple[str, int]]:
        """Split one heading section by token budget while repeating its heading context."""
        heading_prefix = "\n".join(
            f"{'#' * level} {title}" for level, title in heading_path
        )
        if heading_prefix:
            heading_prefix += "\n\n"
        prefix_tokens = self._encoding.encode(heading_prefix)
        if len(prefix_tokens) >= self._chunk_size:
            raise RagConfigurationError("RAG 标题路径超过切片 Token 上限")

        body_tokens = self._encoding.encode(section_content)
        available_body_tokens = self._chunk_size - len(prefix_tokens)
        overlap = min(self._overlap_tokens, available_body_tokens - 1)
        step = available_body_tokens - overlap
        result: list[tuple[str, int]] = []
        for part_index, start in enumerate(range(0, len(body_tokens), step)):
            body = self._encoding.decode(
                body_tokens[start : start + available_body_tokens]
            ).strip()
            if body:
                result.append((f"{heading_prefix}{body}".strip(), part_index))
        return result
